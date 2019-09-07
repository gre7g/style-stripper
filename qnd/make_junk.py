from docx import Document
from docx.styles.style import _ParagraphStyle
from typing import Dict


class DocX(object):
    def __init__(self, template: str) -> None:
        self.doc = Document(template)
        self.style_dict: Dict[str, _ParagraphStyle] = {style.name: style for style in self.doc.styles}
        print(list(self.style_dict.keys()))

    def add_content(self) -> None:
        p = self.doc.paragraphs[0]
        p.text = 'this is the first paragraph, so it should not be indented.'
        p.style = self.style_dict['First Paragraph']
        p = self.doc.add_paragraph('this is the second paragraph, so it ')
        r = p.add_run('should')
        r.italic = True
        p.add_run(' be indented.')
        p.style = self.style_dict['Normal']

    def save_as(self, filename: str) -> None:
        self.doc.save(filename)


if __name__ == "__main__":
    d = DocX(r"..\style_stripper\docx_templates\5x8+bleed.docx")
    d.add_content()
    d.save_as('temp.docx')
