from docx import Document  # noqa
import logging
import re
from typing import List, Tuple, Optional

from style_stripper.data.constants import CONSTANTS
from style_stripper.data.paragraph import Paragraph, QuestionableTick

try:
    from style_stripper.data.book import Book
except ImportError:
    Book = None

# Constants:
LOG = logging.getLogger(__name__)

# Types:
NumPartsType = int
NumChaptersType = int
NumTheEndType = int
ParagraphIndexType = int
ListParaIndexesType = List[ParagraphIndexType]


class OriginalDocx:
    book: Book
    questionable_ticks: List[QuestionableTick]
    paragraphs: List[Paragraph]
    symbolic_divider_indexes: List[ListParaIndexesType]

    def __init__(self, path: str, book: Book):
        self.book = book
        self.questionable_ticks = []

        self.paragraphs = []
        self.symbolic_divider_indexes = []
        doc = Document(path)

        word_count = 0
        for paragraph in doc.paragraphs:
            paragraph_obj = Paragraph()

            for run in paragraph.runs:
                paragraph_obj.add(run.text, run.italic)
            paragraph_obj.get_word_count()
            word_count += paragraph_obj.word_count

            # must_change.append((offset, "‘" if ask_function(self, offset) else "’"))

            # LOG.debug(paragraph_obj.text)
            self.paragraphs.append(paragraph_obj)

        self.book.author = doc.core_properties.author
        self.book.title = doc.core_properties.title
        self.book.word_count = word_count
        self.book.last_modified = doc.core_properties.modified

    def fix_spaces(self):
        for paragraph in self.paragraphs:
            paragraph.fix_spaces(self.book.config)

    def fix_italic_boundaries(self):
        for paragraph in self.paragraphs:
            paragraph.fix_italic_boundaries()

    def fix_ellipses(self):
        for paragraph in self.paragraphs:
            paragraph.fix_ellipses(self.book.config)

    def fix_quotes_and_dashes(self):
        for paragraph in self.paragraphs:
            paragraph.fix_quotes_and_dashes(self.book.config)

    def fix_ticks(self) -> List[QuestionableTick]:
        self.questionable_ticks = []
        for paragraph in self.paragraphs:
            paragraph.fix_ticks(self.questionable_ticks)
        return self.questionable_ticks

    @staticmethod
    def matches_any(searches: List[re.Pattern], text: str):
        return any(search.search(text) for search in searches)

    def find_heading_candidates(
        self,
    ) -> Tuple[NumPartsType, NumChaptersType, NumTheEndType]:
        part = chapter = end = 0
        for paragraph in self.paragraphs:
            if self.matches_any(CONSTANTS.HEADINGS.SEARCH_PART, paragraph.text):
                part += 1
            elif self.matches_any(CONSTANTS.HEADINGS.SEARCH_CHAPTER, paragraph.text):
                chapter += 1
            elif self.matches_any(CONSTANTS.HEADINGS.SEARCH_THE_END, paragraph.text):
                end += 1

        LOG.debug("parts found: %d", part)
        LOG.debug("chapters found: %d", chapter)
        LOG.debug("the end found: %d", end)
        return part, chapter, end

    def style_part_chapter(self, chapter_style: str):
        """Search for part and chapter headings, add style flags, remove blank paragraphs surrounding them"""
        found_part: Optional[Paragraph] = None
        found_chapter: Optional[Paragraph] = None
        group_of_paragraph_indexes: ListParaIndexesType = []
        index: ParagraphIndexType = 0
        while index < len(self.paragraphs):
            paragraph = self.paragraphs[index]
            if paragraph.text == "":
                group_of_paragraph_indexes.append(index)
            elif self.matches_any(CONSTANTS.HEADINGS.SEARCH_PART, paragraph.text):
                group_of_paragraph_indexes.append(index)
                found_part = paragraph
            elif found_part:
                found_part.style = CONSTANTS.STYLING.NAMES.HEADING1
                self.paragraphs[
                    group_of_paragraph_indexes[0] : group_of_paragraph_indexes[-1] + 1
                ] = [found_part]
                index = group_of_paragraph_indexes[0]
                found_part = None
                group_of_paragraph_indexes = []
            elif self.matches_any(CONSTANTS.HEADINGS.SEARCH_CHAPTER, paragraph.text):
                group_of_paragraph_indexes.append(index)
                found_chapter = paragraph
            elif found_chapter:
                found_chapter.style = chapter_style
                self.paragraphs[
                    group_of_paragraph_indexes[0] : group_of_paragraph_indexes[-1] + 1
                ] = [found_chapter]
                index = group_of_paragraph_indexes[0]
                found_chapter = None
                group_of_paragraph_indexes = []
            else:
                group_of_paragraph_indexes = []

            index += 1

    def style_end(self):
        """Search for the end, add style flags, remove blank paragraphs surrounding them"""
        found_end: Optional[Paragraph] = None
        group: ListParaIndexesType = []
        index: ParagraphIndexType = 0
        while index < len(self.paragraphs):
            paragraph = self.paragraphs[index]
            if paragraph.text == "":
                group.append(index)
            elif self.matches_any(CONSTANTS.HEADINGS.SEARCH_THE_END, paragraph.text):
                group.append(index)
                found_end = paragraph
            elif found_end:
                found_end.style = CONSTANTS.STYLING.NAMES.THE_END
                self.paragraphs[group[0] : group[-1] + 1] = [found_end]
                index = group[0]
                found_end = None
                group = []
            else:
                group = []

            index += 1

    def add_end(self):
        end = Paragraph(self.book.config.headings.the_end)
        end.style = CONSTANTS.STYLING.NAMES.THE_END
        self.paragraphs.append(end)

    def find_divider_candidates(self) -> Tuple[int, int]:
        count_of_symbolic = count_of_blanks = 0
        found_blank = found_symbol = False
        group: ListParaIndexesType = []
        for index, paragraph in enumerate(self.paragraphs):
            if not CONSTANTS.DIVIDER.SEARCH.search(paragraph.text):  # No word content
                group.append(index)
                if paragraph.text == "":
                    found_blank = True
                else:
                    found_symbol = True
            else:
                if found_symbol:
                    self.symbolic_divider_indexes.insert(0, group)
                    count_of_symbolic += 1
                elif found_blank:
                    count_of_blanks += 1
                found_blank = found_symbol = False
                group = []

        LOG.debug("symbolic dividers found: %d", count_of_symbolic)
        LOG.debug("symbolic blanks found: %d", count_of_blanks)

        return count_of_symbolic, count_of_blanks

    def replace_symbolic(self):
        indexes: ListParaIndexesType
        while self.symbolic_divider_indexes:
            indexes = self.symbolic_divider_indexes.pop()
            start = indexes[0]
            stop = indexes[-1] + 1
            if self.book.config.divider.replace_with_new:
                self.paragraphs[start:stop] = [
                    Paragraph(
                        self.book.config.divider.new, CONSTANTS.STYLING.NAMES.DIVIDER
                    )
                ]
            else:
                for index in range(start, stop):
                    self.paragraphs[index].style = CONSTANTS.STYLING.NAMES.DIVIDER

    def remove_blanks(self):
        # Indexes are meaningless once we delete paragraphs
        self.symbolic_divider_indexes = []

        index: ParagraphIndexType = 0
        while index < len(self.paragraphs):
            if self.paragraphs[index].text:
                index += 1
            else:
                del self.paragraphs[index]

    def replace_blanks(self):
        # Indexes are meaningless once we delete paragraphs
        self.symbolic_divider_indexes = []

        in_blanks = False
        index: ParagraphIndexType = 0
        while index < len(self.paragraphs):
            if self.paragraphs[index].text:
                in_blanks = False
            else:
                if in_blanks:
                    if self.book.config.divider.replace_with_new:
                        del self.paragraphs[index]
                        continue
                    else:
                        self.paragraphs[index].style = CONSTANTS.STYLING.NAMES.DIVIDER
                else:
                    if self.book.config.divider.replace_with_new:
                        self.paragraphs[index] = Paragraph(self.book.config.divider.new)
                    self.paragraphs[index].style = CONSTANTS.STYLING.NAMES.DIVIDER
                    in_blanks = True

            index += 1

        # May have stuck a divider at the end by accident
        if self.paragraphs[-1].style == CONSTANTS.STYLING.NAMES.DIVIDER:
            del self.paragraphs[-1]
