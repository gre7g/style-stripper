from docx import Document
from docx.styles.style import _ParagraphStyle  # noqa
from glob import glob
import logging
import os
from typing import Dict, Optional, List

from style_stripper.data.constants import CONSTANTS
from style_stripper.data.template_details import TemplateParameters

# Constants:
LOG = logging.getLogger(__name__)

# Types:
DimensionStringType = str


class Templates:
    templates_by_size: Dict[DimensionStringType, List["Template"]]

    def __init__(self):
        self.templates_by_size = {}
        template_pattern = os.path.join(CONSTANTS.PATHS.TEMPLATES_PATH, "*.docx")
        for path in glob(template_pattern):
            match = CONSTANTS.DOCUMENTS.SEARCH_DIMENSIONS.search(path)
            size: DimensionStringType = '%s" x %s"' % match.groups()
            if size not in self.templates_by_size:
                self.templates_by_size[size] = []
            self.templates_by_size[size].append(Template(path))


class Template(TemplateParameters):
    def __init__(self, path: str):
        super(Template, self).__init__()
        self.doc = Document(path)
        self.load = self.load(self.doc)
        self.bleed = "bleed" in path
        self.add_page_break = self.doc.add_page_break
        self.add_section = self.doc.add_section
        self.style_dict: Dict[str, _ParagraphStyle] = {
            style.name: style for style in self.doc.styles
        }
        self.first_paragraph_of_section = True
        self.first_paragraph_of_file = True
        LOG.debug("styles: %r", list(self.style_dict.keys()))

    def add_content(self, text: str = "", style: Optional[str] = None):
        # First paragraph of the file?
        paragraph = (
            self.doc.paragraphs[0]
            if self.first_paragraph_of_file
            else self.doc.add_paragraph()
        )
        self.first_paragraph_of_file = False
        paragraph.text = ""
        if style:
            paragraph.style = style
            self.first_paragraph_of_section = True
        else:
            name = (
                CONSTANTS.STYLING.NAMES.FIRST_PARAGRAPH
                if self.first_paragraph_of_section
                else CONSTANTS.STYLING.NAMES.NORMAL
            )
            paragraph.style = self.style_dict[name]
            self.first_paragraph_of_section = False

        def add_to_paragraph(string: str, italic: bool = False):
            run = paragraph.add_run(string)
            run.italic = italic

        while text:
            match = CONSTANTS.ITALIC.SEARCH.search(text)
            if match:
                before = text[: match.start()]
                if before:
                    add_to_paragraph(before)
                add_to_paragraph(match.group(1), True)
                text = text[match.end() :]
            else:
                add_to_paragraph(text)
                text = ""

    def save_as(self, filename: str):
        self.doc.save(filename)

    def set_properties(self, author, title):
        self.doc.core_properties.author = author
        self.doc.core_properties.title = title


if __name__ == "__main__":
    d = Template(r"docx_templates\5x8+bleed.docx")
    d.add_content("this is the first paragraph, so it should not be indented.")
    d.add_content("this is the second paragraph, so it ❰should❱ be indented.")
    d.save_as(r"..\temp.docx")
