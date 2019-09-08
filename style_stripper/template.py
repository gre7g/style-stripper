from docx import Document
from docx.styles.style import _ParagraphStyle
import logging
import re
from typing import Dict, Optional

# Constants:
LOG = logging.getLogger(__name__)
SEARCH_ITALIC = re.compile(r"❰(.*?)❱")


class Template(object):
    def __init__(self, path: str) -> None:
        self.doc = Document(path)
        self.add_page_break = self.doc.add_page_break
        self.add_section = self.doc.add_section
        self.style_dict: Dict[str, _ParagraphStyle] = {style.name: style for style in self.doc.styles}
        self.first_paragraph_of_section = True
        self.first_paragraph_of_file = True
        LOG.debug("styles: %r", list(self.style_dict.keys()))

    def add_content(self, text: str = "", style: Optional[str] = None) -> None:
        # First paragraph of the file?
        paragraph = self.doc.paragraphs[0] if self.first_paragraph_of_file else self.doc.add_paragraph()
        self.first_paragraph_of_file = False
        paragraph.text = ""
        if style:
            paragraph.style = style
            self.first_paragraph_of_section = True
        else:
            paragraph.style = self.style_dict["First Paragraph" if self.first_paragraph_of_section else "Normal"]
            self.first_paragraph_of_section = False

        def add_to_paragraph(string: str, italic: bool = False) -> None:
            run = paragraph.add_run(string)
            run.italic = italic

        while text:
            match = SEARCH_ITALIC.search(text)
            if match:
                before = text[:match.start()]
                if before:
                    add_to_paragraph(before)
                add_to_paragraph(match.group(1), True)
                text = text[match.end():]
            else:
                add_to_paragraph(text)
                text = ""

    def save_as(self, filename: str) -> None:
        self.doc.save(filename)


if __name__ == "__main__":
    d = Template(r"docx_templates\5x8+bleed.docx")
    d.add_content('this is the first paragraph, so it should not be indented.')
    d.add_content('this is the second paragraph, so it ❰should❱ be indented.')
    d.save_as(r'..\temp.docx')
