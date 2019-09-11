from docx import Document
import logging
from types import FunctionType
from typing import List, Tuple, Optional

from style_stripper.constants import CONSTANTS
from style_stripper.paragraph import Paragraph

# Constants:
LOG = logging.getLogger(__name__)


class OriginalDocx(object):
    def __init__(self, path: str, ask_function: FunctionType) -> None:
        self.paragraphs: List[Paragraph] = []
        self.symbolic_divider_indexes: List[int] = []
        doc = Document(path)
        for paragraph in doc.paragraphs:
            paragraph_obj = Paragraph()

            for run in paragraph.runs:
                paragraph_obj.add(run.text, run.italic)

            paragraph_obj.fix_spaces()
            paragraph_obj.fix_italic_boundaries()
            paragraph_obj.fix_quotes_and_dashes()
            paragraph_obj.fix_ticks(ask_function)

            LOG.debug(paragraph_obj.text)
            self.paragraphs.append(paragraph_obj)

    def find_divider_candidates(self) -> Tuple[int, int]:
        count_of_symbolic = count_of_blanks = 0
        in_blanks = False
        for index, paragraph in enumerate(self.paragraphs):
            for pattern in CONSTANTS.DIVIDER.SEARCH:
                if pattern.search(paragraph.text):
                    count_of_symbolic += 1
                    self.symbolic_divider_indexes.append(index)
                    in_blanks = False
                    break
            else:
                if paragraph.text == "":
                    if not in_blanks:
                        in_blanks = True
                        count_of_blanks += 1
                else:
                    in_blanks = False

        LOG.debug("symbolic dividers found: %d", count_of_symbolic)
        LOG.debug("symbolic blanks found: %d", count_of_blanks)

        return count_of_symbolic, count_of_blanks

    def replace_symbolic(self) -> None:
        for index in self.symbolic_divider_indexes:
            if CONSTANTS.DIVIDER.REPLACE_WITH_NEW:
                self.paragraphs[index] = Paragraph(CONSTANTS.DIVIDER.NEW)
            self.paragraphs[index].style = CONSTANTS.STYLING.NAMES.DIVIDER

    def remove_blanks(self) -> None:
        # Indexes are meaninless once we delete paragraphs
        self.symbolic_divider_indexes = []

        index = 0
        while index < len(self.paragraphs):
            if self.paragraphs[index].text:
                index += 1
            else:
                del self.paragraphs[index]

    def replace_blanks(self) -> None:
        # Indexes are meaninless once we delete paragraphs
        self.symbolic_divider_indexes = []

        in_blanks = False
        index = 0
        while index < len(self.paragraphs):
            if self.paragraphs[index].text:
                index += 1
                in_blanks = False
            else:
                if in_blanks:
                    if CONSTANTS.DIVIDER.REPLACE_WITH_NEW:
                        del self.paragraphs[index]
                    else:
                        self.paragraphs[index].style = CONSTANTS.STYLING.NAMES.DIVIDER
                        index += 1
                else:
                    if CONSTANTS.DIVIDER.REPLACE_WITH_NEW:
                        self.paragraphs[index] = Paragraph(CONSTANTS.DIVIDER.NEW)
                    self.paragraphs[index].style = CONSTANTS.STYLING.NAMES.DIVIDER
                    index += 1
                    in_blanks = True

    def find_heading_candidates(self) -> Tuple[int, int, int]:
        part = chapter = end = 0
        for paragraph in self.paragraphs:
            for pattern in CONSTANTS.HEADINGS.SEARCH_PART:
                if pattern.search(paragraph.text):
                    part += 1
                    break
            for pattern in CONSTANTS.HEADINGS.SEARCH_CHAPTER:
                if pattern.search(paragraph.text):
                    chapter += 1
                    break
            for pattern in CONSTANTS.HEADINGS.SEARCH_THE_END:
                if pattern.search(paragraph.text):
                    end += 1
                    break

        LOG.debug("parts found: %d", part)
        LOG.debug("chapters found: %d", chapter)
        LOG.debug("the end found: %d", end)
        return part, chapter, end

    def style_headings(self, part: Optional[str] = None, chapter: Optional[str] = None, end: Optional[str] = None):
        LOG.debug("styling parts as: %r", part)
        LOG.debug("styling chapters as: %r", chapter)
        LOG.debug("styling the end as: %r", end)

        for paragraph in self.paragraphs:
            for pattern in CONSTANTS.HEADINGS.SEARCH_PART:
                if pattern.search(paragraph.text):
                    paragraph.style = part
                    break
            for pattern in CONSTANTS.HEADINGS.SEARCH_CHAPTER:
                if pattern.search(paragraph.text):
                    paragraph.style = chapter
                    break
            for pattern in CONSTANTS.HEADINGS.SEARCH_THE_END:
                if pattern.search(paragraph.text):
                    paragraph.style = end
                    break

    def remove_dividers_before_headings(self) -> None:
        if not CONSTANTS.DIVIDER.REMOVE_DIVIDERS_BEFORE_HEADINGS:
            return

        # Indexes are meaninless once we delete paragraphs
        self.symbolic_divider_indexes = []

        index = 0
        while index < len(self.paragraphs):
            if self.paragraphs[index].style == CONSTANTS.STYLING.NAMES.DIVIDER:
                if self.paragraphs[index + 1].style in \
                        [CONSTANTS.STYLING.NAMES.HEADING1, CONSTANTS.STYLING.NAMES.HEADING2]:
                    del self.paragraphs[index]
                else:
                    index += 1
            else:
                index += 1
